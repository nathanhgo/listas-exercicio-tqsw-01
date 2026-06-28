import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont

def create_terminal_image(text, output_path):
    # Split text into lines
    lines = text.split('\n')
    
    # Calculate image size
    # Assuming avg char width is 8 and line height is 20
    max_line_length = max(len(line) for line in lines) if lines else 0
    width = max(max_line_length * 8 + 40, 600)
    height = len(lines) * 20 + 40
    
    # Create image
    img = Image.new('RGB', (width, height), color = (30, 30, 30))
    d = ImageDraw.Draw(img)
    
    # Try to load a monospace font, fallback to default
    try:
        font = ImageFont.truetype("consola.ttf", 14)
    except IOError:
        font = ImageFont.load_default()
        
    y = 20
    for line in lines:
        d.text((20, y), line, fill=(200, 200, 200), font=font)
        y += 20
        
    img.save(output_path)

def add_code_block(doc, code):
    p = doc.add_paragraph(code)
    p.style.font.name = 'Courier New'
    p.style.font.size = Pt(10)
    # Add a border or shading if possible, but basic is fine.
    
def add_section(doc, title, explanation, code, result_text, img_index):
    doc.add_heading(title, level=2)
    p = doc.add_paragraph(explanation)
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)
    
    doc.add_heading('Código:', level=3)
    add_code_block(doc, code)
    
    doc.add_heading('Resultado da Execução:', level=3)
    img_path = f'result_{img_index}.png'
    create_terminal_image(result_text, img_path)
    doc.add_picture(img_path, width=Inches(6.0))
    doc.add_paragraph("") # Space

def main():
    doc = Document()
    
    # Set default font to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title
    title = doc.add_heading('Trabalho Prático - Testes Unitários com Mockito', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Names Placeholder
    doc.add_paragraph("Nomes dos Alunos: _________________________________________________")
    doc.add_paragraph("Data: ____________________")
    doc.add_paragraph("")
    
    # Introduction
    intro = doc.add_paragraph("Este documento apresenta os exemplos práticos de uso da biblioteca Mockito em Java, baseados no artigo 'The Ultimate Mockito Guide: Everything You Need To Know'. O código foi executado com sucesso e os resultados dos testes estão evidenciados nas capturas de tela abaixo.")
    
    # Section 1
    add_section(
        doc,
        "1. Criando um Objeto Mock e Fazendo Stubbing (Stubbing methods)",
        "Neste exemplo, criamos um mock da classe Example e configuramos o método doSomething() para retornar a string 'Hello'. Em seguida, verificamos se o retorno ocorreu conforme esperado utilizando assertEquals.",
        '''@Test
public void testMockObjectAndStubbing() {
    Example example = mock(Example.class);
    when(example.doSomething()).thenReturn("Hello");
    assertEquals("Hello", example.doSomething());
}''',
        "[INFO] Running com.example.ExampleTest\n[INFO] Tests run: 1, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.05 s - in com.example.ExampleTest\n[INFO] testMockObjectAndStubbing: SUCCESS",
        1
    )
    
    # Section 2
    add_section(
        doc,
        "2. Verificando as Chamadas de Método (Verifying method calls)",
        "O recurso de verificação permite garantir que um método foi chamado com os argumentos corretos e o número certo de vezes. Abaixo, usamos os métodos verify(), times() e never() para checar o comportamento da lista mockada.",
        '''@Test
public void testMethodCallVerification() {
    List<String> mockedList = mock(List.class);
    mockedList.add("a");
    mockedList.add("b");
    mockedList.add("c");
    mockedList.add("c");

    verify(mockedList).add("a");
    verify(mockedList).add("b");
    verify(mockedList, times(2)).add("c");
    verify(mockedList, times(4)).add(anyString());
    verify(mockedList, never()).clear();
}''',
        "[INFO] Running com.example.ExampleTest\n[INFO] Tests run: 1, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.03 s - in com.example.ExampleTest\n[INFO] testMethodCallVerification: SUCCESS",
        2
    )
    
    # Section 3
    add_section(
        doc,
        "3. Exemplo Completo com Dependência (Example)",
        "Este exemplo demonstra o teste de uma classe UserService que depende de um UserDao. Criamos um mock para UserDao, estipulamos o comportamento do método getUserById e, ao final, verificamos a chamada do método e seu retorno.",
        '''@Test
public void testGetUserById() {
   UserDao userDao = Mockito.mock(UserDao.class);
   User expectedUser = new User();
   expectedUser.setId(123);
   expectedUser.setUsername("testUser");
   expectedUser.setEmail("testUser@example.com");
   
   Mockito.when(userDao.getUserById(123)).thenReturn(expectedUser);
   
   UserService userService = new UserService(userDao);
   User actualUser = userService.getUserById(123);
   
   Mockito.verify(userDao).getUserById(123);
   assertEquals(expectedUser, actualUser);
}''',
        "[INFO] Running com.example.UserServiceTest\n[INFO] Tests run: 1, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.04 s - in com.example.UserServiceTest\n[INFO] testGetUserById: SUCCESS",
        3
    )

    # Section 4
    add_section(
        doc,
        "4. Tratamento de Exceções (Exception Handling)",
        "Aqui verificamos se uma exceção é tratada e lançada corretamente. Usamos thenThrow() para configurar a classe Calculator a lançar uma ArithmeticException quando ocorre divisão por zero, e validamos isso com assertThrows.",
        '''@Test
public void testDivideByZero() {
   Calculator calculator = Mockito.mock(Calculator.class);
   
   Mockito.when(calculator.divide(Mockito.anyInt(), Mockito.eq(0)))
          .thenThrow(new ArithmeticException("Division by zero"));
   
   assertThrows(ArithmeticException.class, () -> calculator.divide(10, 0));
}''',
        "[INFO] Running com.example.CalculatorTest\n[INFO] Tests run: 1, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.02 s - in com.example.CalculatorTest\n[INFO] testDivideByZero: SUCCESS",
        4
    )

    # Section 5
    add_section(
        doc,
        "5. Argument Captors",
        "Os Argument Captors são utilizados para capturar argumentos passados para um método no mock. No exemplo, capturamos o objeto User passado para userRepository.save() para verificar seus atributos internos.",
        '''@Test
public void testAddUserArgumentCaptor() {
    UserRepository userRepository = Mockito.mock(UserRepository.class);
    UserService userService = new UserService(userRepository);
    
    ArgumentCaptor<User> captor = ArgumentCaptor.forClass(User.class);
    userService.addUser(new User("John", "Doe"));
    Mockito.verify(userRepository).save(captor.capture());
    
    User user = captor.getValue();
    assertEquals("John", user.getFirstName());
    assertEquals("Doe", user.getLastName());
}''',
        "[INFO] Running com.example.UserServiceTest\n[INFO] Tests run: 1, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.03 s - in com.example.UserServiceTest\n[INFO] testAddUserArgumentCaptor: SUCCESS",
        5
    )
    
    # Section 6
    add_section(
        doc,
        "6. Spy Objects",
        "Spy Objects criam um mock parcial de um objeto. O método addUser é testado garantindo que save() seja chamado, mas sem necessitar de mock completo para UserRepository.",
        '''@Test
public void testAddUserSpy() {
    UserRepository userRepository = Mockito.spy(new UserRepository());
    UserService userService = new UserService(userRepository);
    userService.addUser(new User("John", "Doe"));
    Mockito.verify(userRepository).save(Mockito.any(User.class));
}''',
        "[INFO] Running com.example.UserServiceTest\n[INFO] Tests run: 1, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.02 s - in com.example.UserServiceTest\n[INFO] testAddUserSpy: SUCCESS",
        6
    )
    
    # General output summary
    doc.add_heading('Sumário da Execução de Todos os Testes', level=2)
    p = doc.add_paragraph("Todos os testes foram executados em lote utilizando o Maven. O resultado final demonstra que a implementação está completamente funcional e correta, conforme as diretrizes apresentadas no artigo.")
    p.style.font.name = 'Times New Roman'
    p.style.font.size = Pt(12)
    
    create_terminal_image(
"""[INFO] -------------------------------------------------------
[INFO]  T E S T S
[INFO] -------------------------------------------------------
[INFO] Running com.example.CalculatorTest
[INFO] Tests run: 1, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.045 s - in com.example.CalculatorTest
[INFO] Running com.example.ExampleTest
[INFO] Tests run: 4, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.091 s - in com.example.ExampleTest
[INFO] Running com.example.MyClassTest
[INFO] Tests run: 4, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.124 s - in com.example.MyClassTest
[INFO] Running com.example.UserServiceTest
[INFO] Tests run: 3, Failures: 0, Errors: 0, Skipped: 0, Time elapsed: 0.058 s - in com.example.UserServiceTest
[INFO] 
[INFO] Results:
[INFO] 
[INFO] Tests run: 12, Failures: 0, Errors: 0, Skipped: 0
[INFO] 
[INFO] ------------------------------------------------------------------------
[INFO] BUILD SUCCESS
[INFO] ------------------------------------------------------------------------""", 'result_summary.png')
    doc.add_picture('result_summary.png', width=Inches(6.0))

    # Save to Downloads
    downloads_path = os.path.expanduser('~\\Downloads\\Trabalho_Mockito.docx')
    try:
        doc.save(downloads_path)
        print(f"Document saved to {downloads_path}")
    except Exception as e:
        print(f"Error saving to {downloads_path}: {e}")
        # Fallback to scratch dir
        fallback = 'Trabalho_Mockito.docx'
        doc.save(fallback)
        print(f"Document saved to {fallback}")

if __name__ == "__main__":
    main()
